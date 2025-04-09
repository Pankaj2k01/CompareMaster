from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 

# Create your views here.
# Home
def home(request):
    return render(request, 'pc/index.html') 

# Web search (Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        percent, link = main.findSimilarity(request.POST['q'])
        percent = round(percent, 2)
    print("Output.....................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Web search file (.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # Creating a pdf file object 
        pdfFileObj = open(request.FILES['docfile'], 'rb') 

        # Creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # Printing number of pages in pdf file 
        print(pdfReader.numPages) 

        # Creating a page object 
        pageObj = pdfReader.getPage(0) 

        # Extracting text from page 
        print(pageObj.extractText()) 

        # Closing the pdf file object 
        pdfFileObj.close() 

    percent, link = main.findSimilarity(value)
    print("Output...................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

# Two text compare (Text)
def twofiletest1(request):
    print("Submitted text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'], request.POST['q2'])
    result = round(result, 2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})
    
# Two text compare (.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1, value2)
    
    print("Output..................!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})

# List Compare
def listCompare(request):
    return render(request, 'pc/list_compare.html')

# Image Compare
def imageCompare(request):
    return render(request, 'pc/image_compare.html')
