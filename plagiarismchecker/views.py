from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 
from django.core.files.storage import FileSystemStorage
from PIL import Image
import numpy as np
from skimage.metrics import structural_similarity as ssim
import cv2


# Create your views here.
# Home
def home(request):
    return render(request, 'pc/index.html') 

def plagiarism(request):
    return render(request, 'pc/plagiarism_checker.html') 

# Web search (Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        try:
            percent, link = main.findSimilarity(request.POST['q'])
            percent = round(percent, 2)
            print("Output.....................!!!!!!!!", percent, link)
            return render(request, 'pc/plagiarism_checker.html', 
                        {'link': link, 'percent': percent})
        except Exception as e:
            print(f"Error during similarity check: {str(e)}")
            return render(request, 'pc/plagiarism_checker.html',
                        {'error': 'Service temporarily unavailable due to high demand. Please try again later.'})

# Web search file (.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    try:
        if str(request.FILES['docfile']).endswith(".txt"):
            value = str(request.FILES['docfile'].read())

        elif str(request.FILES['docfile']).endswith(".docx"):
            document = Document(request.FILES['docfile'])
            for para in document.paragraphs:
                value += para.text

        elif str(request.FILES['docfile']).endswith(".pdf"):
            pdfFileObj = open(request.FILES['docfile'], 'rb') 
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 
            print(pdfReader.numPages) 
            pageObj = pdfReader.getPage(0) 
            print(pageObj.extractText()) 
            pdfFileObj.close() 

        percent, link = main.findSimilarity(value)
        print("Output...................!!!!!!!!", percent, link)
        return render(request, 'pc/plagiarism_checker.html', 
                    {'link': link, 'percent': percent})
    except Exception as e:
        print(f"Error during file processing: {str(e)}")
        return render(request, 'pc/plagiarism_checker.html',
                    {'error': 'Could not process your request. Please try again later.'})

# Text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

# Two text compare (Text)
def twofiletest1(request):
    print("Submitted text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'], request.POST['q2'])
    if result == 0.0:
        return render(request, 'pc/doc_compare.html', {'result': 'No similarity found.'})
    result = round(result, 2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})
    
# Two text compare (.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1, value2)
    
    print("Output..................!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})


def listcompare(request):
    return render(request, 'pc/list_compare.html') 


def imagecompare(request):
    return render(request, 'pc/image_compare.html')

def compare_images(request):
    if request.method == 'POST' and 'image1' in request.FILES and 'image2' in request.FILES:
        image1 = request.FILES['image1']
        image2 = request.FILES['image2']

        # Save the uploaded images
        fs = FileSystemStorage()
        img1_path = fs.save(image1.name, image1)
        img2_path = fs.save(image2.name, image2)

        # Open the images using PIL
        img1 = Image.open(fs.path(img1_path)).convert('L')  # Convert to grayscale
        img2 = Image.open(fs.path(img2_path)).convert('L')  # Convert to grayscale

        # Convert images to numpy arrays
        img1_array = np.array(img1)
        img2_array = np.array(img2)

        # Check if the images have the same dimensions
        if img1_array.shape == img2_array.shape:
            # Calculate Structural Similarity Index (SSIM)
            similarity, diff = ssim(img1_array, img2_array, full=True)
            similarity_percentage = round(similarity * 100, 2)

            # ✅ Fix: If images are 100% similar, do not generate a difference image
            if similarity_percentage == 100.0:
                comparison_result = f"Images are {similarity_percentage}% similar. No differences detected."
                return render(request, 'pc/image_compare.html', {
                    'result': comparison_result
                })

            # If images are not identical, generate the difference image
            diff = (diff * 255).astype("uint8")
            diff_image = Image.fromarray(diff)

            # Convert to in-memory file
            from io import BytesIO
            diff_io = BytesIO()
            diff_image.save(diff_io, format='PNG')
            diff_io.seek(0)

            # Save using FileSystemStorage
            diff_image_path = fs.save("difference_image.png", diff_io)

            comparison_result = f"Images are {similarity_percentage}% similar."
            return render(request, 'pc/image_compare.html', {
                'result': comparison_result,
                'diff_image_url': fs.url(diff_image_path)
            })

        else:
            comparison_result = "Images are of different sizes and cannot be compared."
            return render(request, 'pc/image_compare.html', {'result': comparison_result})

    return render(request, 'pc/image_compare.html', {'result': 'No images uploaded.'})
